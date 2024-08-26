import os
import logging
from flask import Flask, request, redirect, send_from_directory, send_file, render_template, jsonify
import pytesseract
from pdf2image import convert_from_path
from PIL import Image
from docx import Document
import pdfplumber
import io
from dotenv import load_dotenv

# Carregar variáveis do .env
load_dotenv()

# Configurações
UPLOAD_FOLDER = os.getenv('UPLOAD_FOLDER', 'uploads')
RESULT_FOLDER = os.getenv('RESULT_FOLDER', 'results')
pytesseract.pytesseract.tesseract_cmd = os.getenv('TESSERACT_CMD', 'tesseract')
poppler_path = os.getenv('POPPLER_PATH', '/usr/bin/poppler')

# Configuração de logging
logging.basicConfig(filename='app.log', level=logging.INFO)

app = Flask(__name__)

app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['RESULT_FOLDER'] = RESULT_FOLDER
ALLOWED_EXTENSIONS = {'pdf'}

# Funções auxiliares
def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def verificar_pdf_contem_texto(file_path):
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            if page.extract_text():
                return True
    return False

def aplicar_ocr_pdf(caminho_arquivo):
    try:
        logging.info(f"Processando o arquivo: {caminho_arquivo}")
        paginas = convert_from_path(caminho_arquivo, 300, poppler_path=poppler_path)
        texto_completo = ""
        for i, pagina in enumerate(paginas):
            texto = pytesseract.image_to_string(pagina, lang='por')
            logging.info(f"Texto extraído da página {i+1}: {texto[:100]}...")
            texto_completo += texto + "\n\n"
        return texto_completo
    except pytesseract.TesseractError as e:
        logging.error(f"Erro de OCR no arquivo {caminho_arquivo}: {str(e)}")
        raise
    except Exception as e:
        logging.error(f"Erro inesperado: {str(e)}")
        raise

def salvar_como_docx(texto, caminho_saida):
    try:
        doc = Document()
        doc.add_paragraph(texto)
        doc.save(caminho_saida)
    except Exception as e:
        logging.error(f"Erro ao salvar DOCX: {str(e)}")
        raise

def convert_pdf_to_docx(pdf_bytes):
    try:
        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            doc = Document()
            for page in pdf.pages:
                text = page.extract_text()
                if text:
                    doc.add_paragraph(text)
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            return doc_io
    except Exception as e:
        logging.error(f"Erro na conversão do PDF para DOCX: {str(e)}")
        raise

# Rotas do Flask
@app.route('/')
def home():
    return render_template('upload.html')

@app.route('/upload', methods=['POST'])
def upload_file():
    if 'file' not in request.files:
        return redirect(request.url)
    
    file = request.files['file']
    if file.filename == '':
        return redirect(request.url)
    
    if file and allowed_file(file.filename):
        filename = 'uploaded.pdf'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        # Verifica se o PDF contém texto
        if verificar_pdf_contem_texto(file_path):
            logging.info(f"O PDF contém texto, iniciando conversão para DOCX.")
            with open(file_path, 'rb') as f:
                pdf_bytes = f.read()
            docx_file = convert_pdf_to_docx(pdf_bytes)
            return send_file(
                docx_file,
                as_attachment=True,
                download_name=file.filename.replace('.pdf', '.docx')
            )
        else:
            logging.info(f"O PDF não contém texto, iniciando OCR.")
            try:
                texto = aplicar_ocr_pdf(file_path)
                docx_path = os.path.join(app.config['RESULT_FOLDER'], 'resultado.docx')
                salvar_como_docx(texto, docx_path)
                return send_from_directory(app.config['RESULT_FOLDER'], 'resultado.docx', as_attachment=True)
            except Exception as e:
                logging.error(f"Erro durante o processamento: {str(e)}")
                return jsonify({"message": "Erro durante o processamento!"}), 500

    return redirect(request.url)

@app.route('/convert', methods=['POST'])
def convert():
    file = request.files.get('file')
    
    if file and allowed_file(file.filename):
        filename = 'uploaded.pdf'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)

        try:
            if verificar_pdf_contem_texto(file_path):
                with open(file_path, 'rb') as f:
                    pdf_bytes = f.read()
                docx_file = convert_pdf_to_docx(pdf_bytes)
                return send_file(
                    docx_file,
                    as_attachment=True,
                    download_name=file.filename.replace('.pdf', '.docx')
                )
            else:
                texto = aplicar_ocr_pdf(file_path)
                docx_path = os.path.join(app.config['RESULT_FOLDER'], 'resultado.docx')
                salvar_como_docx(texto, docx_path)
                return send_from_directory(app.config['RESULT_FOLDER'], 'resultado.docx', as_attachment=True)
        except Exception as e:
            logging.error(f"Erro durante o processamento: {str(e)}")
            return jsonify({"message": "Erro durante o processamento!"}), 500

    return jsonify({"message": "Falha no upload ou arquivo inválido!"})

if __name__ == '__main__':
    if not os.path.exists(UPLOAD_FOLDER):
        os.makedirs(UPLOAD_FOLDER)
    if not os.path.exists(RESULT_FOLDER):
        os.makedirs(RESULT_FOLDER)
    app.run(host='0.0.0.0', port=5000)
